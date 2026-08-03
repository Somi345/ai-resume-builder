import os
import re
from xhtml2pdf import pisa
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6') # Thin line (1/8 pt per unit)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '444444') # Dark gray border as requested
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_pdf(html_content):
    """
    Uses xhtml2pdf to generate an ATS-friendly PDF from HTML.
    """
    output_path = "temp_resume.pdf"
    
    # We wrap the HTML content in a basic structure if it's not already
    if "<html" not in html_content:
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{ margin: 0; }}
                body {{ font-family: Calibri, Helvetica, Arial, sans-serif; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
    try:
        with open(output_path, "w+b") as result_file:
            pisa_status = pisa.CreatePDF(
                html_content,
                dest=result_file
            )
        
        if pisa_status.err:
            raise Exception("PDF Generation Error")
            
        return output_path
    except Exception as e:
        print(f"PDF Generation Error: {e}")
        raise e

def categorize_skills(skills_list):
    categories = {
        'Programming Languages': ['java', 'python', 'c', 'c++', 'c#', 'cpp', 'javascript', 'typescript', 'go', 'ruby', 'swift', 'kotlin', 'php', 'rust'],
        'Web Technologies': ['html', 'css', 'react', 'react.js', 'reactjs', 'vue', 'angular', 'next.js', 'html5', 'css3', 'tailwind', 'bootstrap'],
        'Backend': ['node.js', 'nodejs', 'express', 'django', 'flask', 'spring boot', 'spring'],
        'Databases': ['mysql', 'postgresql', 'postgres', 'mongodb', 'mongo', 'oracle', 'sql server', 'redis', 'firebase', 'sqlite', 'nosql', 'sql'],
        'Core CS': ['dsa', 'dbms', 'oop', 'operating systems', 'os', 'computer networks', 'cn', 'data structures', 'algorithms'],
        'Cloud & DevOps': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'k8s', 'github actions', 'jenkins', 'ci/cd', 'terraform', 'linux'],
        'Tools': ['git', 'github', 'vs code', 'vscode', 'postman', 'intellij', 'eclipse', 'jira', 'figma', 'webpack', 'npm', 'yarn'],
        'AI/ML': ['tensorflow', 'pytorch', 'opencv', 'scikit-learn', 'pandas', 'numpy', 'keras', 'machine learning', 'deep learning', 'nlp', 'artificial intelligence']
    }
    
    grouped = {k: [] for k in categories.keys()}
    grouped['Other Skills'] = []
    
    for skill_obj in skills_list:
        skill = skill_obj.get('name', '') if isinstance(skill_obj, dict) else skill_obj
        if not skill: continue
        normalized = skill.lower().strip()
        found = False
        for cat, keywords in categories.items():
            if keywords.count(normalized) > 0:
                grouped[cat].append(skill)
                found = True
                break
        if not found:
            grouped['Other Skills'].append(skill)
            
    return {k: v for k, v in grouped.items() if v}

def generate_docx(resume_data):
    """
    Uses python-docx to generate an editable Word document perfectly mirroring the smart rendering.
    """
    output_path = "temp_resume.docx"
    doc = docx.Document()
    
    scale = int(resume_data.get('scaleLevel', 0))
    margin_top = 0.145
    margin_side = 0.187
    font_size = 11.0
    h_before = 8
    if scale == 1:
        margin_top = 0.140; margin_side = 0.180; font_size = 10.5
    elif scale == 2:
        margin_top = 0.135; margin_side = 0.170; font_size = 10.0; h_before = 6
    elif scale == 3:
        margin_top = 0.130; margin_side = 0.160; font_size = 9.5; h_before = 6
    elif scale >= 4:
        margin_top = 0.125; margin_side = 0.150; font_size = 9.5; h_before = 4

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(margin_top)
        section.bottom_margin = Inches(margin_top)
        section.left_margin = Inches(margin_side)
        section.right_margin = Inches(margin_side)

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(0, 0, 0)

    # Helper: Add hyperlink
    def add_hyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(12)
    name_run = name_p.add_run(resume_data.get('fullName', 'YOUR NAME').upper())
    name_run.bold = True
    
    # Scale name
    name_size = 22
    if scale == 1: name_size = 20
    elif scale == 2: name_size = 18
    elif scale == 3: name_size = 16
    elif scale >= 4: name_size = 14
    name_run.font.size = Pt(name_size)
    name_p.paragraph_format.space_after = Pt(0)

    if resume_data.get('targetJob'):
        job_p = doc.add_paragraph()
        job_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        job_run = job_p.add_run(resume_data.get('targetJob'))
        job_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        
        job_size = 11
        if scale == 1: job_size = 10.5
        elif scale == 2: job_size = 10
        elif scale >= 3: job_size = 9.5
        job_run.font.size = Pt(job_size)
        job_p.paragraph_format.space_after = Pt(1)

    # Combined contact line
    contact1 = []
    contact2 = []
    if resume_data.get('email'): contact1.append(resume_data.get('email'))
    if resume_data.get('phone'): contact1.append(resume_data.get('phone'))
    if resume_data.get('location'): contact1.append(resume_data.get('location'))
    
    if resume_data.get('github'): contact2.append(("GitHub", resume_data.get('github')))
    if resume_data.get('linkedin'): contact2.append(("LinkedIn", resume_data.get('linkedin')))
    if resume_data.get('portfolio'): contact2.append(("Portfolio", resume_data.get('portfolio')))

    contactAll = contact1 + contact2
    if contactAll:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(2)
        
        for idx, item in enumerate(contactAll):
            if idx > 0: cp.add_run(" | ")
            if isinstance(item, tuple):
                add_hyperlink(cp, item[0], item[1])
            else:
                cp.add_run(item)

    # Section Helper
    def add_section_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(h_before)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(font_size)
        add_bottom_border(p)

    def add_item_header(left_text, right_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        
        # Setup tabs for right alignment
        tab_stops = p.paragraph_format.tab_stops
        try:
            tab_stops.add_tab_stop(Inches(8.27 - 2*margin_side), WD_TAB_ALIGNMENT.RIGHT)
        except Exception:
            pass
            
        run_left = p.add_run(left_text)
        run_left.bold = True
        if right_text:
            p.add_run('\t' + right_text)

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.25)
        p.text = text

    # Summary
    summary = resume_data.get('summary', '').strip()
    if summary:
        add_section_heading('OBJECTIVE')
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Education
    if resume_data.get('education'):
        add_section_heading('EDUCATION')
        for edu in resume_data.get('education', []):
            add_item_header(edu.get('degree', 'Degree'), edu.get('year', ''))
            
            school = edu.get('school', 'College')
            raw_grade = edu.get('grade', '')
            grade_str = raw_grade if ('%' in raw_grade or 'cgpa' in raw_grade.lower() or 'percentage' in raw_grade.lower()) else (f"CGPA: {raw_grade}" if raw_grade else "")
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            tab_stops = p.paragraph_format.tab_stops
            try:
                tab_stops.add_tab_stop(Inches(8.27 - 2*margin_side), WD_TAB_ALIGNMENT.RIGHT)
            except Exception:
                pass
            
            p.add_run(school).italic = True
            if grade_str:
                p.add_run('\t' + grade_str)
                
        # Coursework below education
        if resume_data.get('coursework'):
            cw = [c.get('name') if isinstance(c, dict) else c for c in resume_data.get('coursework', [])]
            cwList = ', '.join([c for c in cw if c])
            if cwList:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(0)
                run_cw = p.add_run('Relevant Coursework: ')
                run_cw.bold = True
                p.add_run(cwList)
    
    # Skills
    if resume_data.get('skills'):
        add_section_heading('TECHNICAL SKILLS')
        categorized = categorize_skills(resume_data.get('skills', []))
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False
        # Set exact column widths matching 20% and 80% approx on A4 page (8.27 inches wide)
        usable_width = 8.27 - 2*margin_side
        table.columns[0].width = Inches(usable_width * 0.20)
        table.columns[1].width = Inches(usable_width * 0.80)
        
        for cat, items in categorized.items():
            row_cells = table.add_row().cells
            p_cat = row_cells[0].paragraphs[0]
            p_cat.paragraph_format.space_after = Pt(0)
            p_cat.add_run(cat).bold = True
            
            p_items = row_cells[1].paragraphs[0]
            p_items.paragraph_format.space_after = Pt(0)
            p_items.add_run(', '.join(items))

    # Projects
    if resume_data.get('projects'):
        add_section_heading('PROJECTS')
        for proj in resume_data.get('projects', []):
            add_item_header(proj.get('title', 'Project Name'), proj.get('technologies', ''))
            
            desc = proj.get('description', '')
            if desc:
                bullets = [b.strip() for b in desc.split('\n') if b.strip()]
                for b in bullets[:2]:
                    text = re.sub(r'^- ', '', b).strip()
                    if text:
                        words = text.split()
                        if len(words) > 20: text = ' '.join(words[:20]) + '...'
                        add_bullet(text)

    # Experience
    if resume_data.get('experience'):
        add_section_heading('EXPERIENCE')
        for exp in resume_data.get('experience', []):
            dates = f"{exp.get('startDate', '')} - {exp.get('endDate', '')}" if exp.get('endDate') else exp.get('startDate', '')
            add_item_header(exp.get('title', 'Role'), dates)
            
            if exp.get('company'):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.add_run(exp.get('company'))
                
            desc = exp.get('description', '')
            if desc:
                bullets = [b.strip() for b in desc.split('\n') if b.strip()]
                for b in bullets[:2]:
                    text = re.sub(r'^- ', '', b).strip()
                    if text:
                        words = text.split()
                        if len(words) > 20: text = ' '.join(words[:20]) + '...'
                        add_bullet(text)

    # Internships
    if resume_data.get('internships'):
        add_section_heading('INTERNSHIPS')
        for intern in resume_data.get('internships', []):
            dates = f"{intern.get('startDate', '')} - {intern.get('endDate', '')}" if intern.get('endDate') else intern.get('startDate', '')
            add_item_header(intern.get('title', 'Role'), dates)
            
            if intern.get('company'):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.add_run(intern.get('company'))
                
            desc = intern.get('description', '')
            if desc:
                bullets = [b.strip() for b in desc.split('\n') if b.strip()]
                for b in bullets[:2]:
                    text = re.sub(r'^- ', '', b).strip()
                    if text:
                        words = text.split()
                        if len(words) > 20: text = ' '.join(words[:20]) + '...'
                        add_bullet(text)

    # Certifications
    if resume_data.get('certifications'):
        add_section_heading('CERTIFICATIONS')
        for cert in resume_data.get('certifications', [])[:1]:
            if isinstance(cert, dict):
                line = cert.get('name', '')
                if cert.get('organization'): line += f", {cert.get('organization')}"
                if cert.get('year'): line += f", {cert.get('year')}"
                if line: add_bullet(line)

    # Coding Profiles
    if resume_data.get('codingProfiles'):
        valid_profiles = [p for p in resume_data.get('codingProfiles', []) if p.get('link') and p.get('link').strip() != '']
        if valid_profiles:
            add_section_heading('CODING PROFILES')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            for i, p_data in enumerate(valid_profiles):
                if i > 0:
                    p.add_run(" | ")
                add_hyperlink(p, p_data.get('platform', p_data.get('name', 'Profile')), p_data.get('link'))

    # Achievements
    if resume_data.get('achievements'):
        add_section_heading('ACHIEVEMENTS')
        for ach in resume_data.get('achievements', [])[:1]:
            name = ach.get('name', '') if isinstance(ach, dict) else ach
            if name:
                add_bullet(name)

    # Languages
    if resume_data.get('languages'):
        valid_langs = []
        for l in resume_data.get('languages', [])[:2]:
            if isinstance(l, dict):
                name = l.get('name', '')
                if name: valid_langs.append(name)
            elif l:
                valid_langs.append(str(l))
        if valid_langs:
            add_section_heading('LANGUAGES')
            p = doc.add_paragraph(' | '.join(valid_langs))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)

    doc.save(output_path)
    return output_path
